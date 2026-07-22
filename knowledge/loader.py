"""
乐意AI - 文档加载器

支持加载 PDF、TXT、Word、Markdown 格式的文档。
"""

import os
from typing import List, Optional
from dataclasses import dataclass


@dataclass
class Document:
    """文档对象"""
    content: str
    metadata: dict

    def __repr__(self) -> str:
        return f"Document(content={len(self.content)}chars, metadata={self.metadata})"


def load_document(file_path: str) -> Optional[Document]:
    """
    加载单个文档

    Args:
        file_path: 文件路径

    Returns:
        Document 对象，失败返回 None
    """
    if not os.path.exists(file_path):
        return None

    ext = os.path.splitext(file_path)[1].lower()
    filename = os.path.basename(file_path)
    metadata = {"source": filename, "path": file_path}

    try:
        if ext == ".txt":
            return _load_txt(file_path, metadata)
        elif ext == ".md":
            return _load_md(file_path, metadata)
        elif ext == ".pdf":
            return _load_pdf(file_path, metadata)
        elif ext == ".docx":
            return _load_docx(file_path, metadata)
        else:
            metadata["error"] = f"不支持的文件格式: {ext}"
            return Document(content="", metadata=metadata)
    except Exception as e:
        metadata["error"] = str(e)
        return Document(content="", metadata=metadata)


def load_documents(file_paths: List[str]) -> List[Document]:
    """
    批量加载文档

    Args:
        file_paths: 文件路径列表

    Returns:
        Document 列表
    """
    docs = []
    for path in file_paths:
        doc = load_document(path)
        if doc and doc.content:
            docs.append(doc)
    return docs


def _load_txt(file_path: str, metadata: dict) -> Document:
    """加载TXT文件"""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    metadata["format"] = "txt"
    return Document(content=content, metadata=metadata)


def _load_md(file_path: str, metadata: dict) -> Document:
    """加载Markdown文件（去除格式标记，保留纯文本）"""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    metadata["format"] = "markdown"
    return Document(content=content, metadata=metadata)


def _load_pdf(file_path: str, metadata: dict) -> Document:
    """加载PDF文件"""
    import fitz  # PyMuPDF
    doc = fitz.open(file_path)
    pages = []
    for page_num, page in enumerate(doc, 1):
        text = page.get_text().strip()
        if text:
            pages.append(text)
    doc.close()
    content = "\n\n".join(pages)
    metadata["format"] = "pdf"
    metadata["pages"] = len(pages)
    return Document(content=content, metadata=metadata)


def _load_docx(file_path: str, metadata: dict) -> Document:
    """加载Word文档"""
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    content = "\n".join(paragraphs)
    metadata["format"] = "docx"
    return Document(content=content, metadata=metadata)


def list_documents(doc_dir: str) -> List[str]:
    """列出文档目录中的支持文件"""
    supported = {".txt", ".md", ".pdf", ".docx"}
    files = []
    if os.path.exists(doc_dir):
        for f in os.listdir(doc_dir):
            ext = os.path.splitext(f)[1].lower()
            if ext in supported:
                files.append(os.path.join(doc_dir, f))
    return files